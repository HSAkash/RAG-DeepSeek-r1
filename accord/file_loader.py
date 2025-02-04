from dataclasses import dataclass
from pathlib import Path
from streamlit.runtime.uploaded_file_manager import UploadedFile
from pypdfium2 import PdfDocument
from docx import Document
from accord.utils import get_config


config = get_config()

TEXT_FILE_EXTENSION = ".txt"
PDF_EXTENSION = ".pdf"
MD_FILE_EXTENSION = ".md"
DOCX_EXTENSION = ".docx"



@dataclass
class File:
    name: str
    content: str

def extract_pdf_content(file_path: Path) -> str:
    """
    Extract text content from pdf file
    Args:
        file_path (Path): Path to the pdf file
    Returns:
    str: Text content of the pdf file
    """
    pdf = PdfDocument(file_path)
    content = ""
    for page in pdf:
        # content += page.get_text()
        text_page = page.get_textpage()
        content += f"{text_page.get_text_bounded()}"
    return content

def extract_docx_content(file_path: Path) -> str:
    """
    Extract text content from docx file
    Args:
        file_path (Path): Path to the docx file
    Returns:
    str: Text content of the docx file
    """
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text


def load_file(file_path: UploadedFile) -> File:
    """
    Load file content based on file extension
    Args:
        file (UploadedFile): File uploaded by user
    Raises:
        ValueError: If file extension is not allowed
    Returns:
    File: Loaded file content
    """
    file_extension = Path(file_path.name).suffix
    file_name = file_path.name


    content = ""
    if file_extension not in config.documentUpload.ALLOWED_FILE_EXTENSIONS:
        raise ValueError(f"File extension {file_extension} not allowed")

    if file_extension == TEXT_FILE_EXTENSION or file_extension == MD_FILE_EXTENSION:
        content = file_path.getvalue().decode("utf-8")
    elif file_extension == PDF_EXTENSION:
        content = extract_pdf_content(file_path)
    elif file_extension == DOCX_EXTENSION:
        content = extract_docx_content(file_path)
    return File(name=file_name, content=content)